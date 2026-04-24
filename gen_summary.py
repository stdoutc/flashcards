# -*- coding: utf-8 -*-
"""
Internship summary generator for Meng Jiaxu.
Usage: d:\\anaconda3\\python.exe gen_summary.py
Output: docs/summary/internship-summary-mengjx.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 图片路径解析 ──────────────────────────────────────────────────────────────
BASE_DIR   = os.path.dirname(os.path.abspath(__file__))
IMG_DIR    = os.path.join(BASE_DIR, 'docs', 'img')
GIF_DIR    = os.path.join(BASE_DIR, 'docs', 'gif')
OUT_PATH   = os.path.join(BASE_DIR, 'docs', '总结', '实习总结-孟嘉许.docx')

def find_img(prefix):
    """按前缀匹配 IMG_DIR 中的图片文件，返回完整路径；找不到返回 None。"""
    for f in os.listdir(IMG_DIR):
        if f.startswith(prefix):
            return os.path.join(IMG_DIR, f)
    return None

def find_gif(prefix):
    """按前缀匹配 GIF_DIR 中的 gif 文件，返回完整路径；找不到返回 None。"""
    for f in os.listdir(GIF_DIR):
        if f.startswith(prefix):
            return os.path.join(GIF_DIR, f)
    return None

IMG_31  = find_img('3.1')   # 卡组管理
IMG_32  = find_img('3.2')   # 编辑卡片
IMG_33  = find_img('3.3')   # 图片插入显示
IMG_41  = find_img('4.1')   # 学习复习
IMG_42  = find_img('4.2')   # 在学n张
IMG_51  = find_img('5.1')   # 统计面板
IMG_61  = find_img('6.1')   # 导入导出备份
IMG_62  = find_img('6.2')   # 导入备份
IMG_63  = find_img('6.3')   # 数据备份

# ── 文档工具函数 ──────────────────────────────────────────────────────────────
def set_font(run, size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, size=15, bold=True)
    return p

def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, size=13, bold=True)
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing      = Pt(22)
    p.paragraph_format.space_after       = Pt(4)
    run = p.add_run(text)
    set_font(run, size=12)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Pt(28)
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run('• ' + text)
    set_font(run, size=12)
    return p

def add_caption(doc, text):
    """图片说明文字，居中。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_font(run, size=10, color=(100, 100, 100))
    return p

def add_image(doc, path, width_cm=14, caption=None):
    """插入图片，若文件不存在则跳过。"""
    if path and os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Cm(width_cm))
        if caption:
            add_caption(doc, caption)
    else:
        if path:
            print(f'[WARN] 图片不存在，跳过：{path}')

def add_two_images(doc, path1, path2, caption=None):
    """在同一段中并排插入两张图片。"""
    paths = [p for p in [path1, path2] if p and os.path.exists(p)]
    if not paths:
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for path in paths:
        run = para.add_run()
        run.add_picture(path, width=Cm(6.5))
        run.add_text('   ')
    if caption:
        add_caption(doc, caption)

def add_table(doc, headers, rows, caption=None):
    """插入带表头的表格。"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # 表头
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                set_font(run, size=11, bold=True)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 数据行
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, cell_text in enumerate(row_data):
            row_cells[ci].text = cell_text
            for para in row_cells[ci].paragraphs:
                for run in para.runs:
                    set_font(run, size=11)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        add_caption(doc, caption)
    return table

# ── 正文开始 ──────────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.2)
    section.right_margin  = Cm(3.2)

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(14)
run = title.add_run('实  习  总  结')
set_font(run, size=18, bold=True)

# ═══════════════════════════════════════════════════════════════
# 一、实习单位及项目背景
# ═══════════════════════════════════════════════════════════════
add_heading1(doc, '一、实习单位及项目背景')

add_body(doc,
    '本次 Co-op 实习依托学校 Co-op 项目平台开展，实习时间为 2025 年 1 月 2 日至 '
    '2025 年 4 月 30 日，历时近四个月。我与曾钰同学共同组队，在导师指导下自主立项，'
    '完成了从需求分析、系统设计、开发实现到测试部署的完整软件工程流程。项目以"AI 赋能"'
    '为核心理念，聚焦解决大学生日常学习中普遍存在的问题：知识点记了就忘、复习缺乏计划、笔记零散难以整理。'
)

add_body(doc,
    '项目最终成果是一款名为"联想记忆卡片"的多功能学习工具，定位为面向学习者的 AI 赋能'
    '记忆应用。其核心思路是用卡片来承载知识点，用间隔重复算法来规划复习，结合 AI 智能'
    '制卡与联想图谱等功能帮助记忆。应用已完成 Web 端与 iOS 移动端的开发和部署，Web 端'
    '在线地址为 https://rococo-beignet-799198.netlify.app/，源代码托管于 GitHub'
    '（https://github.com/stdoutc/flashcards）。'
)

# ═══════════════════════════════════════════════════════════════
# 二、实习工作内容
# ═══════════════════════════════════════════════════════════════
add_heading1(doc, '二、实习工作内容')

add_body(doc,
    '在本次项目中，我主要承担四个方面的工作：应用初期构思与系统架构设计、UI 设计与'
    '用户体验优化、核心功能开发实现，以及版本管理与应用部署。以下分模块说明具体工作内容。'
)

# ── 2.1 ──
add_heading2(doc, '（一）应用构思与系统架构设计')

add_body(doc,
    '项目启动阶段，我主导了整体功能规划和技术选型。在调研 Anki 等同类产品的基础上，'
    '确定了以"卡组-卡片"为核心数据模型的信息架构，并规划了基础功能（卡片管理、学习复习、'
    '统计面板、数据导入导出）与两大特色功能（AI 识别制卡、联想图谱）。'
)

add_body(doc,
    '技术选型上，考虑到跨平台需求，我选择了 React + TypeScript 作为整体框架，使用 Vite '
    '进行构建，采用 React Router 管理路由，移动端通过 Expo 实现。这一架构使 Web 端与 iOS '
    '端能够共用同一套业务逻辑代码，仅在表现层做平台差异化处理，显著降低了维护成本。'
    '下表展示了本项目的完整技术栈：'
)

add_table(doc,
    headers=['层次', '技术 / 工具', '用途'],
    rows=[
        ['整体框架',   'React + TypeScript', '组件化 UI 与类型安全'],
        ['构建工具',   'Vite',               '快速开发构建与热更新'],
        ['路由管理',   'React Router',       '多页面 SPA 路由'],
        ['卡片渲染',   'Markdown + LaTeX',   '富文本与数学公式显示'],
        ['移动端',     'Expo',               'iOS 跨平台编译与预览'],
        ['版本管理',   'Git / GitHub',       '代码托管与协作'],
        ['Web 部署',   'Netlify',            '持续集成与自动部署'],
    ],
    caption='表 1  项目技术栈一览'
)

# ── 2.2 ──
add_heading2(doc, '（二）基础功能开发')

add_body(doc,
    '基础功能包括卡组与卡片管理、学习与复习流程、统计面板以及数据导入导出，'
    '是整个应用的核心骨架。'
)

add_body(doc,
    '卡组是存储卡片的容器，支持新建、重命名、删除，以及单独设置每日新卡上限。卡片内容'
    '支持纯文本、Markdown（加粗、斜体、代码块、表格）及 LaTeX 数学公式，同时支持插入本地'
    '图片。这一设计突破了传统闪卡的纯文本限制，适合理工科知识点的记录与复习。'
)

add_two_images(doc, IMG_31, IMG_32, caption='图 1  卡组管理界面（左）与卡片编辑界面（右）')
add_image(doc, IMG_33, width_cm=9, caption='图 2  图片插入与预览显示效果')

add_body(doc,
    '学习与复习流程中，我实现了参考 Anki SM-2 算法的间隔重复调度系统。新卡受每日新卡'
    '上限约束，待复习的旧卡到期后自动加入复习队列。用户对每张卡片按掌握程度打分，'
    '系统据此动态调整卡片的难度系数和下次复习间隔，实现个性化记忆曲线。'
    '下表展示了四档评级对应的算法参数：'
)

add_table(doc,
    headers=['评级', '难度系数变化', '下次复习间隔'],
    rows=[
        ['艰难', '−0.20', '10 分钟（重新进入学习步）'],
        ['困难', '−0.15', '上次间隔 × 1.2'],
        ['普通', '不变',  '上次间隔 × 难度系数'],
        ['简单', '+0.15', '上次间隔 × 难度系数 × 1.3'],
    ],
    caption='表 2  间隔重复调度算法评级参数'
)

add_body(doc,
    '此外，应用还实现了"在学 n 张"功能：当天新卡学完后，可临时突破每日上限继续学习，'
    '第二天上限自动恢复，方便用户在考前冲刺时灵活使用。'
)

add_two_images(doc, IMG_41, IMG_42, caption='图 3  学习复习界面（左）与在学 n 张功能（右）')

add_body(doc,
    '统计面板提供了综合学习情况、每日学习数量、近 14 天学习量趋势、掌握度分布及各卡组'
    '状态等多维度数据可视化，帮助用户直观了解自身学习进展。'
)

add_image(doc, IMG_51, width_cm=13, caption='图 4  统计面板界面')

add_body(doc,
    '数据导入导出功能支持三种操作，如下表所示：'
)

add_table(doc,
    headers=['操作', '说明', '格式'],
    rows=[
        ['单卡组导出', '分享或备份某一个卡组',             'JSON'],
        ['全部数据导出备份', '跨设备迁移时一次性导出所有数据', 'JSON'],
        ['导入',      '覆盖现有数据，操作前需确认',        'JSON'],
    ],
    caption='表 3  导入导出功能说明'
)

add_image(doc, IMG_61, width_cm=13, caption='图 5  导入导出备份界面')

# ── 2.3 ──
add_heading2(doc, '（三）特色功能开发')

add_body(doc,
    '在完成基础功能的基础上，我与曾钰同学共同开发了两项特色功能：AI 识别制卡与联想图谱。'
)

add_body(doc,
    'AI 识别制卡功能针对一个普遍存在的问题而设计：课堂上用手机拍下的笔记和题目，'
    '事后还需手动逐条录入整理，耗时且容易搁置。该功能对接了国产（豆包）'
    '与国际主流（ChatGPT、Gemini）AI 模型接口，用户上传课堂截图或手写笔记图片后，AI '
    '自动识别其中的知识点并生成结构化闪卡，实现"图片→整理→闪卡→归档"的全流程自动化。'
    '开发过程中，我通过多轮测试对比了不同模型的响应速度与稳定性，并持续优化 Prompt 提示词'
    '以提升识别准确率。同时提供"极速模式"与"精确模式"两档供用户按需选择。'
)

add_body(doc,
    '联想图谱功能的设计灵感来源于"知识网络化"学习理念——英语单词变形、数学概念关联、'
    '专业课程逻辑梳理等场景都需要在知识点之间建立有意义的联结，而传统闪卡只能孤立地记忆'
    '单个知识点。为此，我实现了以闪卡为节点的树状联想图：用户可在卡组内搜索卡片并'
    '将其组织为父子关系的树形结构，系统提供可视化图谱预览、节点拖拽编辑、缩略图导航，'
    '以及基于图谱路径的专项学习流程，使零散知识点形成结构化体系，构建过程即为深度梳理。'
)

# ═══════════════════════════════════════════════════════════════
# 三、工作思考与收获
# ═══════════════════════════════════════════════════════════════
add_heading1(doc, '三、工作思考与收获')

add_heading2(doc, '（一）AI 辅助编程工具的实践认识')

add_body(doc,
    '本次实习中，我大量使用了 Cursor AI Agent 等 AI 编程工具辅助开发。AI 工具在加速'
    '原型开发、处理重复性代码等方面确实省了不少时间，但它对整体架构的把握能力有限，'
    '生成的代码有时存在逻辑漏洞或与项目风格不一致，仍需逐一核查和修改。'
)

add_body(doc,
    '这段经历让我对 AI 编程工具有了更理性的认识——它更像是一个执行力强但需要人来把方向'
    '的助手。架构怎么划分、模块之间如何解耦、出了问题从哪里入手排查，这些判断还是得靠'
    '自己积累下来的工程经验，AI 替代不了。'
)

add_heading2(doc, '（二）前端工程与跨平台开发能力的积累')

add_body(doc,
    '通过本次实习，我完整实践了基于 Node.js 的前端开发流程，从 npm 包管理、Vite 构建'
    '配置到 TypeScript 类型系统，工具链的每个环节都是边踩坑边摸清楚的。这也是我第一次'
    '把 React 用在一个有实际功能的项目里，从组件拆分、状态管理到路由配置，走了一遍之后'
    '对前端开发的整体脉络清晰了很多。'
)

add_body(doc,
    '在移动端开发方面，通过 Expo 实现 iOS 端应用时，我对"共享业务逻辑、差异化表现层"'
    '的跨平台思路有了直观的体会。调度算法（scheduler.ts）和联想树（assocTree.ts）作为'
    '纯逻辑模块与 UI 完全分离，Web 端和移动端都可以直接复用，后期修改时只需动一处，'
    '省去了不少来回同步的麻烦。'
)

add_heading2(doc, '（三）软件架构思维的建立')

add_body(doc,
    '在实际开发中，我逐渐养成了将代码分层组织的习惯，把表现层（UI 组件与页面）、'
    '业务逻辑层（调度算法、联想树、数据处理）和数据访问层（本地存储）明确区分开来。'
    '这样做的好处在开发后期体现得很明显——比如需要调整复习算法时，只改 scheduler.ts'
    '就够了，完全不用动页面代码，改起来放心很多。'
)

add_body(doc,
    '项目初期对卡片数据模型（Card、Deck、ReviewLog）的设计也让我意识到前期结构定得好'
    '有多重要。后来陆续加入统计面板和联想图谱这两个功能时，基本上都能顺着现有结构扩展，'
    '没有出现需要推翻重来的情况。'
)

add_heading2(doc, '（四）团队协作能力的提升')

add_body(doc,
    '本次实习是我第一次在真实项目里用 Git + GitHub 做多人协作。分支管理、代码合并、'
    '冲突解决这些操作一开始并不顺手，但用多了也就熟悉了。分工上我主要负责功能实现，'
    '曾钰同学负责测试调试和反馈问题，这种搭配在实践中效果不错——她发现的不少问题是我'
    '自己测试时很难注意到的，沟通一两句就能定位到原因。做完这个项目之后，'
    '我对"两个人配合"和"一个人单打独斗"之间的差别有了比较直接的感受。'
)

add_heading2(doc, '（五）创新意识的培养')

add_body(doc,
    '联想图谱和 AI 识别制卡这两个功能都不在最初的需求列表里，是开发过程中想到并加进去的。'
    '联想图谱的想法来自于我自己学习时的感受——单张卡片背孤立的知识点效果有限，把相关的'
    '内容串起来记忆会好很多。AI 识别制卡则是观察到课堂拍的照片往往就躺在相册里没人整理，'
    '想着能不能直接从图片生成卡片，省掉手动录入这一步。两个功能做出来之后反馈都还不错，'
    '这让我意识到，有用的新功能往往来自对自身使用习惯的反思，而不是凭空构想。'
)

# ═══════════════════════════════════════════════════════════════
# 四、未来规划
# ═══════════════════════════════════════════════════════════════
add_heading1(doc, '四、未来规划与期望')

add_body(doc,
    '本次实习结束后，五一假期过后我将正式回到课堂。我计划认真投入专业课的学习，努力夯实'
    '电气专业的理论基础，将本次实习中积累的工程实践经验与课堂知识相互印证，补足理论短板，'
    '提升专业综合能力。'
)

add_body(doc,
    '在下一阶段的实习规划上，我的目标是争取申请到科技类公司的技术岗实习，将专业所学'
    '真正带入实际工程场景中加以检验和深化，为未来的职业发展积累更扎实的基础。'
)

# ── 保存 ──────────────────────────────────────────────────────────────────────
doc.save(OUT_PATH)
print('Saved: ' + OUT_PATH)

# 字数统计
total_chars = sum(len(p.text) for p in doc.paragraphs)
print('Total chars: ' + str(total_chars))
