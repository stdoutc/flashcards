# -*- coding: utf-8 -*-
"""
将实习总结内容填入鉴定表指定方框。
Usage: d:\\anaconda3\\python.exe fill_form.py
Output: docs/总结/附件4.Co-op实习总结鉴定表（新）- 已填.docx
"""

import copy, os
from docx import Document
from docx.oxml.ns import qn
import lxml.etree as etree

BASE_DIR    = os.path.dirname(os.path.abspath(__file__))
SUMMARY     = os.path.join(BASE_DIR, 'docs', '总结', '实习总结-孟嘉许.docx')
FORM_IN     = os.path.join(BASE_DIR, 'docs', '总结', '附件4.Co-op实习总结鉴定表（新）.docx')
FORM_OUT    = os.path.join(BASE_DIR, 'docs', '总结', '附件4.Co-op实习总结鉴定表（新）- 已填.docx')

# ── 图片关系迁移工具 ──────────────────────────────────────────────────────────

def copy_image_rels(src_doc, dst_doc, elem):
    """
    遍历 elem 中所有图片引用（a:blip r:embed / v:imagedata r:id），
    将对应图片数据从 src_doc 复制到 dst_doc，并把 elem 中的关系 ID 替换为新 ID。
    返回修改后的 elem（已 deepcopy）。
    """
    elem = copy.deepcopy(elem)

    BLIP_TAG  = qn('a:blip')
    EMBED_ATT = qn('r:embed')
    LINK_ATT  = qn('r:link')

    for blip in elem.iter(BLIP_TAG):
        for att in (EMBED_ATT, LINK_ATT):
            old_rid = blip.get(att)
            if not old_rid:
                continue
            # 找到源文档中该 rid 对应的图片 part
            try:
                src_part = src_doc.part.rels[old_rid].target_part
            except KeyError:
                continue
            # 将图片 part 加入目标文档，获取新 rid
            new_rid = dst_doc.part.relate_to(
                src_part,
                src_doc.part.rels[old_rid].reltype,
            )
            blip.set(att, new_rid)

    return elem


def copy_elem(src_doc, dst_doc, elem):
    """深拷贝元素，同时迁移图片关系。"""
    tag = etree.QName(elem.tag).localname
    if tag in ('drawing', 'pict'):
        return copy_image_rels(src_doc, dst_doc, elem)
    # 段落/表格：先深拷贝，再递归处理子元素中的图片
    new_elem = copy.deepcopy(elem)
    for drawing in new_elem.iter(qn('w:drawing')):
        parent = drawing.getparent()
        idx    = list(parent).index(drawing)
        fixed  = copy_image_rels(src_doc, dst_doc, drawing)
        parent.remove(drawing)
        parent.insert(idx, fixed)
    return new_elem


# ── 主逻辑 ───────────────────────────────────────────────────────────────────

src_doc  = Document(SUMMARY)
form_doc = Document(FORM_IN)

table    = form_doc.tables[0]
target   = table.cell(5, 0)          # 实习总结方框（gridSpan=8）
tc       = target._tc

# 1. 保留第一段（"实 习 总 结" 标题行），清除其余占位段落
keep_first = tc[0]                   # 第一个 <w:p>（标题）
# 删除 tc 中所有 w:p 和 w:tbl，保留 w:tcPr
to_remove = [
    child for child in list(tc)
    if etree.QName(child.tag).localname in ('p', 'tbl')
]
for child in to_remove:
    tc.remove(child)

# 重新插入标题段
tc.append(keep_first)

# 2. 将总结文档 body 中的所有顶层元素逐一复制进目标单元格
body = src_doc.element.body
for child in body:
    tag = etree.QName(child.tag).localname
    if tag in ('p', 'tbl'):
        new_elem = copy_elem(src_doc, form_doc, child)
        tc.append(new_elem)
    # sectPr 等其他标签跳过

# 3. 为所有嵌入表格显式补全线框（TableGrid 样式在目标文档中可能不存在）
BORDER_XML = (
    '<w:tblBorders'
    ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:top    w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:left   w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:right  w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '</w:tblBorders>'
)
CELL_BORDER_XML = (
    '<w:tcBorders'
    ' xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:top    w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:left   w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '<w:right  w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    '</w:tcBorders>'
)

def set_table_borders(tbl):
    """为表格及其所有单元格显式设置线框。"""
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = etree.SubElement(tbl._tbl, qn('w:tblPr'))
    # 替换或新增 tblBorders
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(etree.fromstring(BORDER_XML))
    # 每个单元格也显式设置边框
    for row in tbl.rows:
        for cell in row.cells:
            tcPr = cell._tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = etree.SubElement(cell._tc, qn('w:tcPr'))
            old_cb = tcPr.find(qn('w:tcBorders'))
            if old_cb is not None:
                tcPr.remove(old_cb)
            tcPr.append(etree.fromstring(CELL_BORDER_XML))

for tbl in target.tables:
    set_table_borders(tbl)

# 4. 确保单元格末尾有一个空段落（Word 规范要求）
last = tc[-1]
if etree.QName(last.tag).localname != 'p':
    from lxml import etree as _e
    tc.append(_e.fromstring('<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

# 5. 保存
form_doc.save(FORM_OUT)
print('Saved: ' + FORM_OUT)
